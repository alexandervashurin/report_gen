import os

from docx.shared import Inches
from docx import Document
from docxtpl import DocxTemplate
import datetime
from report_app.report_users.modules_users.employees_and_departments import STAFF
from report_app.report_users.modules_users.employees_and_departments import DEPARTMENTS
from report_app.report_users.modules_users.employees_and_departments import POSITIONS
import report_app.report_users.modules_users.equipment_staff as eq
from report_app.views import directory

date_time = str(datetime.datetime.now().date())
d_time = date_time.split(
    '-')[2] + '.' + date_time.split('-')[1] + '.' + date_time.split('-')[0] \
         + ' г.'


def function_report(results):
    document = Document()
    run = document.add_paragraph().add_run()
    font = run.font
    font.bold = True
    table = document.add_table(rows=3, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Пользователь'
    hdr_cells[1].text = str(dict(STAFF).get(int(results[0])))
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'Должность'
    hdr_cells[1].text = dict(POSITIONS).get(int(results[1]))
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Подразделение'
    hdr_cells[1].text = dict(DEPARTMENTS).get(int(results[2]))

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Номер п/п'
    hdr_cells[1].text = 'Наименование устройства'
    hdr_cells[2].text = 'Характеристика оборудования'
    hdr_cells[3].text = 'Инв.№'
    hdr_cells[4].text = 'Телефон'
    hdr_cells[5].text = 'Подпись'

    row_cells = table.add_row().cells
    row_cells[0].text = '1'
    row_cells[1].text = dict(eq.TYPES_EQUIPMENTS).get(int(results[3]))
    row_cells[2].text = dict(eq.TYPES_PROCESSORS).get(int(results[4]))
    if results[5] is not None:
        row_cells[3].text = str(results[5])
    else:
        row_cells[3].text = ''

    row_cells[5].text = str(results[8])

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = ''
    hdr_cells[2].text = ''
    hdr_cells[3].text = ''
    hdr_cells[4].text = ''
    hdr_cells[5].text = ''

    row_cells = table.add_row().cells
    row_cells[0].text = '2'
    row_cells[1].text = 'Марка монитора: ' + \
        dict(eq.TYPES_EQUIPMENTS).get(int(results[6]))
    row_cells[2].text = 'Видеокарта: ' + \
        dict(eq.TYPE_VIDEOS).get(int(results[7]))
    row_cells[3].text = ''
    row_cells[4].text = ''
    row_cells[5].text = ''

    document.add_page_break()
    document.save(str(directory) + "/" + "отчёты" + "/" + "demo.docx")

    return document


def function_report2(results):
    doc = DocxTemplate(str(directory) + "/" + "отчёты" + "/" + "demo.docx")
    inv_OS = dict(eq.INVENTORY_NUMBER_MAIN_EQUIPMENTS).get((results[13]))
    inv_TK = dict(eq.INVENTORY_NUMBER_THIN_CLIENT).get((results[21]))
    inv_nout = dict(eq.INVENTORY_NUMBER_NOTEBOOK).get((results[22]))
    if dict(eq.TYPES_EQUIPMENTS).get((results[3])) == 'Системный блок':
        res = inv_OS
    elif dict(eq.TYPES_EQUIPMENTS).get((results[3])) == 'Тонкий клиент НР ' \
                                                        'Flexible':
        res = inv_TK
    elif dict(eq.TYPES_EQUIPMENTS).get((results[3])) == 'Ноутбук Dell ' \
                                                        'Latitude 3510 15.6 ':
        res = inv_nout
    elif dict(eq.TYPES_EQUIPMENTS).get((results[3])) == 'Ноутбук Dell ' \
                                                        'Inspiron 5490':
        res = inv_nout
    elif dict(eq.TYPES_EQUIPMENTS).get((results[3])) == 'Ноутбук HP ProBook ' \
                                                        '430 G5, 13.3':
        res = inv_nout
    elif dict(eq.TYPES_EQUIPMENTS).get((results[3])) == 'Ноутбук-трансформер ' \
                                                        ' HP Pavilion x360':
        res = inv_nout
    else:
        res = "нет"

    context = {'фио': str(dict(STAFF).get((results[0]))),
               'должность': dict(POSITIONS).get((results[1])),
               'подразделение': dict(DEPARTMENTS).get((results[2])),
               'устройство': dict(eq.TYPES_EQUIPMENTS).get((results[3])),
               'типпроцессора': dict(eq.TYPES_PROCESSORS).get((results[4])),
               'инв_принтера': dict(eq.INVENTORY_NUMBER_PRINTERS).get(
                   (results[5])),
               'марка_монитора': dict(eq.BRAND_MONITORS).get((results[6])),
               'видеокарта': dict(eq.TYPE_VIDEOS).get((results[7])),
               'номер телефона': results[8],
               'ОС': dict(eq.TYPE_OSES).get((results[9])),
               'диагональ_монитора': results[10],
               'инв_монитора': dict(eq.INVENTORY_NUMBER_MONITORS).get(
                   (results[11])),
               'принтер_локальный': dict(eq.AVAILABILITY).get((results[12])),


               'ИБП': dict(eq.AVAILABILITY).get((results[14])),

               'колонки': dict(eq.AVAILABILITY).get((results[15])),
               'инв_колонки': dict(eq.INVENTORY_NUMBER_SOUND_SPEAKERS).get((results[16])),

               'телефон': dict(eq.NAME_PHONES).get((results[17])),
               'inv_num_phone': dict(eq.INVENTORY_NUMBER_PHONES).get((results[18])),
               'number_phone': results[19],

               'тип_монитора': dict(eq.TYPES_MONITORS).get((results[20])),

               'инв_тк':  dict(eq.INVENTORY_NUMBER_THIN_CLIENT).get((results[21])),

               'пр1': dict(eq.TYPES_EQUIPMENTS).get((results[23])),
               'пр2': dict(eq.TYPES_EQUIPMENTS).get((results[24])),




               'инв_пр1': results[25],
               'инв_пр2': results[26],



               'инв_ИБП': dict(eq.UPS).get((results[27])),

               'инв_вебкам': dict(eq.INVENTORY_NUMBER_WEBCAM).get((results[
                   28])),
               'вебкам': dict(eq.AVAILABILITY).get((results[29])),

               'дата': d_time
               }
    doc.render(context)
    doc.save(str(directory) + "/" + "отчёты" + "/" + "report.docx")

    return doc
