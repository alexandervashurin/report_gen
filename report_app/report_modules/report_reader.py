#!/usr/bin/python3
import os
import re
from docx import Document
import datetime

directory = '/home/alex/report_project/report_project/media/documents2'
dt = datetime.datetime.now().date()
temp_d = directory + '/' + str(dt).split('-')[0] + '/' + str(dt).split('-')[1] + '/'
dir_files = os.listdir(directory)
slova = ['По состоянию на', '№', 'ПСН', 'БЦВ', 'СТПР', 'ПЧ', 'УЗАРД', 'кондиц', 'УОВ', 'МПСУиД', 'ИЧМ', 'СВОП',
         'ДПС', 'Неисправна клим', 'КУ', 'Клим', 'требуется устранение утечки', 'камера видеонаблюдения', 'в 1 контуре',
         'во 2 контуре',
         'БУК', 'КЛИМ', 'ССТ', 'БЛОК']


# TODO добавить в виде переменных ключевые слова


def method_docx(dir_y):
    def m_stringos(str_s):
        for i in str_s:
            for j in slova:
                if j in i:
                    doc2.add_paragraph("", style='ListBullet').add_run(i).bold = True

            else:
                pass

    for ii in range(dir_files.__len__() - 1):
        ps = temp_d + dir_y[ii]
        document = Document(ps)

        ll = [i for i in document.paragraphs if i.runs]
        title = [i.text for i in document.paragraphs if i.runs and 'По состоянию' in i.text]

        r = []
        for i in ll:
            r = r + i.runs

        tt = [i for i in r if i.bold]

        stringos = [i.text for i in tt if ' ' in i.text]

        doc2 = Document()
        doc2.add_heading(title, 0)
        tp = []
        m_stringos(str_s=stringos)
        ph = re.compile(r'\d\d.\s\w\w\d\w\s№\d\d\d')  # номер пункта и наименование электровоза
        p1 = re.compile(r'\S')
        tu = [i for i in tp if (p1.findall(i) != [])]
        n = dir_files[ii].strip('.docx')
        # handle.write(bold_str)
        dir_docx = '/home/alex/report_project/report_project/media/documents2'
        doc2.save(dir_docx + '/' + str(dt).split('-')[0] + '/' + str(dt).split('-')[1] + '/' + str(dt) + '.docx')

    return doc2
